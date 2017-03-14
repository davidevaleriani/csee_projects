import cherrypy
from cherrypy.lib import auth_basic
import os
import zipfile
# Install antiword with apt-get install antiword
import subprocess
import shutil
# Install python-docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xlrd
import csv
import glob
from PyPDF2 import PdfFileReader
from io import IOBase
import numpy as np
import sys
from openpyxl import load_workbook

use_antiword = False


class HomePage(object):
    @cherrypy.expose
    def index(self):
        return """
            <html>
            <head>
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/3.3.4/css/bootstrap.min.css">
                <script src="https://maxcdn.bootstrapcdn.com/bootstrap/3.3.4/js/bootstrap.min.js"></script>
                <title>CSEE Abstract Booklet Creator</title>
                <style>
                /* Wrapper for page content to push down footer */
                html, body {
                  padding-top: 0px;
                  font-size: 15px;
                  height: 100%;
                }

                #wrap {
                  min-height: 100%;
                  height: auto !important;
                  height: 100%;
                  /* Negative indent footer by it's height */
                  margin: 0 auto -40px;
                  padding-bottom: 10px;
                }

                /* Set the fixed height of the footer here */
                #push,
                #footer {
                  height: 40px;
                  padding-top: 10px;
                }
                #footer {
                  font-size: 12px;
                  background-color: #ddd;
                  text-align: center;
                }
                </style>
            </head>
            <body>
            <div id="wrap">
                <div class="container">
                    <h2 align="center">Welcome to the Abstract Booklet Creator</h2>
                    <p align="center">This application allows you to create the booklet from a zip of abstracts.</p>
                    <br>
                    <form method="post" action="get_booklet" enctype="multipart/form-data" class="form-horizontal">
                        <div class="form-group">
                            <label for="abstracts" class="col-xs-2 col-xs-offset-2 control-label">Zip file</label>
                            <div class="col-xs-6">
                                <input type="file" name="abstracts" id="abstracts" />
                            </div>
                        </div>
                        <div class="form-group">
                            <div class="col-xs-offset-6 col-sm-6">
                                <button type="submit" class="btn btn-primary">Generate</button>
                            </div>
                        </div>
                    </form>
                </div>
                <div id="push"></div>
            </div>
            <div id="footer">
                <div class="container">
                    <p class="muted credit">Created by <a href="http://www.davidevaleriani.it">Davide Valeriani</a>.</p>
                </div>
            </div>
            </body>
            </html>
            """

    @cherrypy.expose
    def get_booklet(self, abstracts):
        # Get list of students for "backup"
        students = {}
        poster_order = {}
        students_by_surname = {}
        if os.path.exists("CE301LST.csv"):
            # Load list of students from CSV file
            # Full name, first names, degree, regno, surname
            with open('CE301LST.csv', 'rb') as csvfile:
                spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')
                for row in spamreader:
                    # strudents[regno] = (first names, surname, degree)
                    students[row[3]] = (row[1], row[4], row[2])
                    students_by_surname[row[4].upper()] = (row[1], row[2])
        elif os.path.exists("students_list.xlsx"):
            wb = load_workbook(filename='students_list.xlsx', read_only=True)
            student_doc = wb.get_sheet_by_name(wb.get_sheet_names()[0])
            row_iterator = student_doc.iter_rows()
            # Skip the first row
            next(row_iterator)
            first_name_column = 4
            last_name_column = 5
            regno_column = 6
            number_column = 10
            degree_column = 7
            for row in row_iterator:
                if row[regno_column-1].value is not None:
                    students[row[regno_column-1].value] = (row[first_name_column - 1].value,
                                                           row[last_name_column - 1].value,
                                                           row[degree_column - 1].value)
                    poster_order[row[regno_column-1].value] = row[number_column - 1].value
        # Create the directory where the abstract will be temporarily saved
        if not os.path.isdir("tmp"):
            os.mkdir("tmp")
        # Extract files from the zip file
        if isinstance(abstracts, IOBase) and hasattr(abstracts, "file"):
            abstracts = abstracts.file
        #zipf = zipfile.ZipFile(abstracts, 'r')
        #zipf.extractall("tmp/")
        # Init the booklet document
        booklet_doc = Document("booklet_template.docx")
        booklet_doc.add_page_break()
        processed_files = []
        files_not_inserted = []
        files_to_be_checked = []
        files_duplicated = []
        students_added = []
        # Sort abstracts by creation time
        #mtime = lambda f: os.stat(os.path.join('tmp/', f)).st_mtime
        #list_of_abstracts = list(sorted(os.listdir('tmp/'), key=mtime))[::-1]
        order_by_number = lambda f: poster_order[int(f)]-1
        # For each abstract form
        print("TOTAL STUDENTS", len(list(students.keys())))
        for student_regno in sorted(students.keys(), key=order_by_number):
            student_first = students[student_regno][0]
            student_last = students[student_regno][1]
            degree = students[student_regno][2]
            project_title = ""
            supervisor_first = ""
            supervisor_last = ""
            abstract = ""
            if len(glob.glob("tmp/"+str(student_regno)+"*")) == 0:
                print("!! WARNING: missing Document for student", student_regno)
            else:
                f = glob.glob("tmp/"+str(student_regno)+"*")[0]
                # Rename file to unix-like format (remove spaces, etc.)
                os.rename(f, f.replace(" ", "-").lower())
                f = f.replace(" ", "-").lower()
                if f in processed_files:
                    print("Duplicate!")
                    continue
                processed_files.append(f)
                abstract_has_started = False
                if f[-4:] == "docx":
                    # If the abstracts are in .docx format
                    document = Document(f)
                    for p in document.paragraphs:
                        p = p.text
                        if "Title" in p and not project_title:
                            try:
                                project_title = p.split("Title:")[1].lstrip()
                                project_title = ' '.join(project_title.split())
                            except:
                                print("WARNING %s: check TITLE" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                project_title = p
                        elif "Student First Name" in p and not student_first:
                            try:
                                if len(p.split("Student First Name(s):")) > 1:
                                    student_first = p.split("Student First Name(s):")[1].lstrip()
                                elif len(p.split("Student First Name:")) > 1:
                                    student_first = p.split("Student First Name:")[1].lstrip()
                                else:
                                    raise
                                student_first = ' '.join(student_first.split())
                            except:
                                print("WARNING %s: check STUDENT NAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                student_first = p
                        elif "Student Surname" in p and not student_last:
                            try:
                                student_last = p.split("Student Surname:")[1].lstrip()
                                student_last = ' '.join(student_last.split())
                            except:
                                print("WARNING %s: check STUDENT SURNAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                student_last = p
                        elif "Student" in p and not student_first and not student_last:
                            try:
                                student_name = p.split(":")[1].lstrip()
                                student_first = ' '.join(student_name.split(" ")[:-1])
                                student_last = student_name.split(" ")[-1]
                            except:
                                print("WARNING %s: check STUDENT FULL NAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                student_last = p
                        elif "Supervisor First" in p:
                            try:
                                if len(p.split("Supervisor First Name(s):")) > 1:
                                    supervisor_first = p.split("Supervisor First Name(s):")[1].lstrip()
                                elif len(p.split("Supervisor First Name:")) > 1:
                                    supervisor_first = p.split("Supervisor First Name:")[1].lstrip()
                                else:
                                    raise
                                supervisor_first = ' '.join(supervisor_first.split())
                            except:
                                print("WARNING %s: check SUPERVISOR NAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                supervisor_first = p
                        elif "Supervisor Surname" in p:
                            try:
                                supervisor_last = p.split("Supervisor Surname:")[1].lstrip()
                                supervisor_last = ' '.join(supervisor_last.split())
                            except:
                                print("WARNING %s: check SUPERVISOR SURNAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                supervisor_last = p
                        elif "Supervisor" in p and (not supervisor_first or not supervisor_last):
                            try:
                                supervisor = p.split(":")[1].lstrip()
                                supervisor_first = ' '.join(supervisor.split(" ")[:-1])
                                supervisor_last = supervisor.split(" ")[-1]
                            except:
                                print("WARNING %s: check SUPERVISOR FULL NAME" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                student_last = p
                        elif "Degree Course" in p and not degree:
                            try:
                                degree = p.split("Degree Course:")[1].lstrip()
                                degree = ' '.join(degree.split())
                            except:
                                print("WARNING %s: check DEGREE" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                degree = p
                        elif abstract_has_started:
                            abstract += p
                        elif "ABSTRACT" in p.upper():
                            if "Abstract" in p:
                                abstract = p.split("Abstract")[1].lstrip()
                                abstract_has_started = True
                            elif "ABSTRACT" in p:
                                abstract = p.split("ABSTRACT")[1].lstrip()
                                abstract_has_started = True
                            elif "abstract" in p:
                                abstract = p.split("abstract")[1].lstrip()
                                abstract_has_started = True
                            else:
                                print("WARNING %s: check ABSTRACT" % f)
                                if f in processed_files:
                                    processed_files.remove(f)
                                if f not in files_to_be_checked:
                                    files_to_be_checked.append(f)
                                abstract = p
                            if abstract and abstract[0] == ".":
                                abstract = abstract[1:]
                elif f[-3:] == "doc":
                    # If the abstracts are in .doc format
                    try:
                        # Convert the .doc to .txt
                        txt_file = subprocess.check_output(["antiword", "%s" % f], universal_newlines=True)
                    except:
                        print("WARNING: unable to convert", f, "- SKIP")
                        continue
                    counter = 0
                    for field in txt_file.split('\n'):
                        field = field.lstrip()
                        if "Title" in field:
                            project_title = field.split("Title:")[1].lstrip()
                        elif "Student".upper() in field.upper():
                            student = field.split(":")[1].lstrip()
                            student_first = (' '.join(student.split(" ")[:-1]))
                            student_last = student.split(" ")[-1]
                        elif "Supervisor".upper() in field.upper():
                            supervisor = field[len("Supervisor:"):].lstrip()
                            supervisor_first = ' '.join(supervisor.split(":")[:-1])
                            supervisor_last = supervisor.split(":")[-1]
                        elif "Abstract".upper() in field.upper():
                            abstract = field[len("Abstract."):].lstrip()
                            abstract += "\n".join(txt_file.split("\n")[counter+1:])
                            break
                        counter += 1
                    if not degree:
                        reformatted_student_name = (student.split(" ")[-1]+", "+" ".join(student.split(" ")[:-1])).upper()
                        try:
                            degree = students[reformatted_student_name]
                        except:
                            degree = ""
                elif f[-3:].upper() == "PDF":
                    pdf = PdfFileReader(open(f, "rb")).getPage(0).extractText().replace('\n', '').replace('\t', ' ').replace("  ", " ").replace("!", "")
                    try:
                        project_title = pdf.split("itle:")[1].lstrip().split("Student")[0]
                    except:
                        print("WARNING %s: check TITLE" % f)
                        if f in processed_files:
                            processed_files.remove(f)
                        if f not in files_to_be_checked:
                            files_to_be_checked.append(f)
                    if not student_first or not student_last:
                        try:
                            student = pdf.split("Student:")[1].lstrip().split("Supervisor")[0].rstrip()
                            student_first = " ".join(student.split(" ")[:-1])
                            student_last = student.split(" ")[-1]
                            reformatted_student_name = (student.split(" ")[-1] + ", " + " ".join(student.split(" ")[:-1])).upper()
                        except:
                            print("WARNING %s: check STUDENT FULL NAME" % f)
                            if f in processed_files:
                                processed_files.remove(f)
                            if f not in files_to_be_checked:
                                files_to_be_checked.append(f)
                    else:
                        reformatted_student_name = (student_last + ", " + student_first).upper()
                    try:
                        supervisor = pdf.split("Supervisor:")[1].lstrip().split("Abstract")[0].rstrip()
                        supervisor_first = " ".join(supervisor.split(" ")[:-1])
                        supervisor_last = supervisor.split(" ")[-1]
                    except:
                        print("WARNING %s: check SUPERVISOR FULL NAME" % f)
                        if f in processed_files:
                            processed_files.remove(f)
                        if f not in files_to_be_checked:
                            files_to_be_checked.append(f)
                    try:
                        abstract = pdf.split("Abstract")[1].lstrip()
                        if abstract[:2] == ". ":
                            abstract = abstract[2:]
                    except:
                        print("WARNING %s: check ABSTRACT" % f)
                        if f in processed_files:
                            processed_files.remove(f)
                        if f not in files_to_be_checked:
                            files_to_be_checked.append(f)
                    if not degree:
                        try:
                            degree = students[reformatted_student_name]
                        except:
                            degree = ""
            #if not all([project_title, student_first, student_last, supervisor_first, supervisor_last]):
            #    print("WARNING: file %s not inserted" % f)
            #    if f in processed_files:
            #        processed_files.remove(f)
            #    if f not in files_not_inserted:
            #        files_not_inserted.append(f)
                # print("TITLE:", project_title)
                # print("STUDENT:", student_first, " > ", student_last)
                # print("SUPERVISOR:", supervisor_first, " > ",supervisor_last)
                # print("DEGREE COURSE:", degree)
                # print("ABSTRACT:", abstract)
                # print()
            #    continue
            if not degree:
                if student_last.upper() in list(students_by_surname.keys()):
                    degree = students_by_surname[student_last.upper()][1]
                else:
                    print("WARNING: file %s doesn't have a degree course" % f)
                    if f not in files_to_be_checked:
                        files_to_be_checked.append(f)
                    # print("TITLE:", project_title)
                    # print("STUDENT:", student_first, student_last)
                    # print("SUPERVISOR:", supervisor_first, supervisor_last)
                    # print("DEGREE COURSE:", degree)
                    # print("ABSTRACT:", abstract)
                    # print()
            if not abstract:
                print("WARNING: file %s doesn't have an abstract" % f)
                if f not in files_to_be_checked:
                    files_to_be_checked.append(f)
                # print("TITLE:", project_title)
                # print("STUDENT:", student_first, student_last)
                # print("SUPERVISOR:", supervisor_first, supervisor_last)
                # print("DEGREE COURSE:", degree)
                # print("ABSTRACT:", abstract)
                # print()
            if student_regno in students_added:
                print("WARNING: file %s is a duplicate" % f)
                if f in processed_files:
                    processed_files.remove(f)
                if f not in files_duplicated:
                    files_duplicated.append(f)
                continue
            students_added.append(student_regno)
            # Adding to the main document
            heading = booklet_doc.add_paragraph()
            paragraph_format = heading.paragraph_format
            run = heading.add_run("\n"+project_title.title())
            run.bold = True
            font = run.font
            font.size = Pt(13)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = booklet_doc.add_table(3, 3)
            table.allow_autofit = False
            for i in range(3):
                row = table.rows[i]
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "260")
                trHeight.set(qn('w:hRule'), "exact")
                trPr.append(trHeight)
            table.cell(0, 0).paragraphs[0].add_run('Name:').bold = True
            table.cell(0, 1).paragraphs[0].add_run(student_first.title()+" "+student_last.title()).bold = True
            table.cell(1, 0).paragraphs[0].add_run('Degree course:').bold = True
            table.cell(1, 1).paragraphs[0].add_run(degree.title()).bold = True
            table.cell(2, 0).paragraphs[0].add_run('Supervisor:').bold = True
            table.cell(2, 1).paragraphs[0].add_run(supervisor_first.title()+" "+supervisor_last.title()).bold = True
            abstract_paragraph = booklet_doc.add_paragraph()
            run = abstract_paragraph.add_run('\nAbstract.\n')
            run.bold = True
            font = run.font
            font.size = Pt(12)
            run = abstract_paragraph.add_run(abstract)
            run.bold = False
            font = run.font
            font.size = Pt(12)
        # Save the report
        booklet_doc.save("booklet.docx")

        # Remove the marks
        #shutil.rmtree('tmp')

        return """
        <html>
        <body>
        <h2>Files added: %d</h2>
        <h2>Files to be checked: %d</h2>
        <h3><ul>%s</ul><h3>
        <h2>Files not inserted: %d</h2>
        <h3><ul>%s</ul><h3>
        <h2>Files duplicated: %d</h2>
        <h3><ul>%s</ul><h3>
        <h2>Total: %d</h2>
        <h1 align="center"><a href="download_booklet">Download booklet</a></h1>
        </body>
        </html>
        """ % (len(processed_files),
               len(files_to_be_checked),
               ''.join(["<li>"+filename+"</li>" for filename in files_to_be_checked]),
               len(files_not_inserted),
               ''.join(["<li>"+filename+"</li>" for filename in files_not_inserted]),
               len(files_duplicated),
               ''.join(["<li>"+filename+"</li>" for filename in files_duplicated]),
               len(processed_files+files_not_inserted+files_duplicated))

    @cherrypy.expose
    def download_booklet(self):
        # Return the file to download
        booklet_doc = open("booklet.docx", 'r')
        return cherrypy.lib.static.serve_fileobj(booklet_doc, disposition='attachment', name="booklet.docx")

if __name__ == '__main__':
    cherrypy.server.socket_host = '0.0.0.0'
    cherrypy.quickstart(HomePage())
